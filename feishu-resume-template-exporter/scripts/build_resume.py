#!/usr/bin/env python3
"""Build a sample-style Chinese resume from structured Feishu content."""

import argparse
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# This family is present on the macOS machines where the Skill is distributed
# and is recognized by both Word and LibreOffice during PDF conversion.
FONT = "Source Han Serif CN"
BODY_SIZE = 7.2


def set_run_font(run, name=FONT, size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def hide_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"}, bottom={"val": "nil"},
                left={"val": "nil"}, right={"val": "nil"},
            )


def set_paragraph_border_bottom(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def clean_document(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    while len(doc.sections) > 1:
        doc.sections[-1]._sectPr.getparent().remove(doc.sections[-1]._sectPr)


def style_page(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(0.9)
    sec.bottom_margin = Cm(0.85)
    sec.left_margin = Cm(1.15)
    sec.right_margin = Cm(1.15)
    sec.header_distance = Cm(0.6)
    sec.footer_distance = Cm(0.6)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(BODY_SIZE)


def add_text(paragraph, text, size=BODY_SIZE, bold=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_header(doc, data, avatar):
    # Matching left/right rails visually center the identity block on the page
    # while keeping the photo independently right-aligned.
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(14.7)
    table.columns[2].width = Cm(2.0)
    hide_table_borders(table)
    _, identity, photo = table.rows[0].cells
    set_cell_margins(identity, top=0, bottom=0)
    set_cell_margins(photo, top=0, bottom=0)
    # Keep the portrait as an independent right rail, while centering the
    # two-line identity block within the portrait's height.
    identity.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    photo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = identity.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0.5)
    add_text(p, data["name"], size=14.5, bold=True)
    p = identity.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_text(p, data["contact"], size=8.1)
    if avatar:
        p = photo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(avatar), width=Cm(2.0), height=Cm(2.0))


def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    add_text(p, title, size=9.0, bold=True)
    set_paragraph_border_bottom(p)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2 + 0.22 * level)
    p.paragraph_format.first_line_indent = Cm(-0.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.1)
    p.paragraph_format.line_spacing = 0.94
    add_text(p, "• ", size=BODY_SIZE, bold=True)
    add_text(p, text, size=BODY_SIZE)
    return p


def add_entry(doc, entry):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.3)
    table.columns[1].width = Cm(3.4)
    hide_table_borders(table)
    left_cell, right_cell = table.rows[0].cells
    set_cell_margins(left_cell, top=0, bottom=0, start=0, end=0)
    set_cell_margins(right_cell, top=0, bottom=0, start=0, end=0)
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1.2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    left = entry["organization"] + (f"｜{entry['role']}" if entry.get("role") else "")
    add_text(p, left, size=7.9, bold=True)
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(1.2)
    p.paragraph_format.space_after = Pt(0)
    date_run = add_text(p, entry["period"], size=7.2)
    date_run.font.color.rgb = RGBColor(45, 45, 45)
    if entry.get("intro"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0.1)
        p.paragraph_format.keep_with_next = True
        add_text(p, entry["intro"], size=BODY_SIZE)
    for item in entry.get("bullets", []):
        add_bullet(doc, item)


def build(data, template, avatar, output):
    shutil.copy2(template, output)
    doc = Document(output)
    clean_document(doc)
    style_page(doc)
    add_header(doc, data, avatar)
    add_section_title(doc, "个人总结")
    for item in data.get("summary", []):
        add_bullet(doc, item)
    for section in data.get("sections", []):
        if not section.get("entries"):
            continue
        add_section_title(doc, section["title"])
        for entry in section["entries"]:
            add_entry(doc, entry)
    if data.get("education"):
        add_section_title(doc, "教育经历")
        for edu in data["education"]:
            add_entry(doc, {
                "organization": edu["school"], "role": "", "period": edu.get("period", ""),
                "intro": edu.get("detail", ""), "bullets": edu.get("bullets", []),
            })
    if data.get("skills"):
        add_section_title(doc, "精通技能")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_text(p, "、".join(data["skills"]), size=BODY_SIZE)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--avatar", type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    for key in ("name", "contact"):
        if not data.get(key):
            raise ValueError(f"Missing required field: {key}")
    args.output.parent.mkdir(parents=True, exist_ok=True)
    build(data, args.template, args.avatar, args.output)


if __name__ == "__main__":
    main()
